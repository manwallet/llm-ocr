from flask import Flask, render_template, request, jsonify, send_file
import os
import requests
import base64
from PIL import Image
import io
import time
import configparser
import json
import tempfile
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB

class OCRProcessor:
    def __init__(self):
        self.config_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ocr_config.ini')
        self.api_key = ""
        self.api_url = "https://api.openai.com/v1/chat/completions"
        self.api_model = "gpt-4-vision-preview"
        self.max_retries = 3
        self.load_config()
    
    def load_config(self):
        if not os.path.exists(self.config_file):
            return
        
        config = configparser.ConfigParser()
        try:
            config.read(self.config_file)
            if 'API' in config:
                self.api_key = config['API'].get('api_key', "")
                self.api_url = config['API'].get('api_url', self.api_url)
                self.api_model = config['API'].get('api_model', self.api_model)
                self.max_retries = int(config['API'].get('max_retries', self.max_retries))
        except Exception as e:
            print(f"加载配置时出错: {str(e)}")
    
    def save_config(self):
        config = configparser.ConfigParser()
        config['API'] = {
            'api_key': self.api_key,
            'api_url': self.api_url,
            'api_model': self.api_model,
            'max_retries': str(self.max_retries)
        }
        
        try:
            with open(self.config_file, 'w') as f:
                config.write(f)
        except Exception as e:
            print(f"保存配置时出错: {str(e)}")
    
    def process_image(self, image_path):
        max_retries = self.max_retries
        retry_count = 0
        retry_delay = 2
        
        while retry_count <= max_retries:
            try:
                with open(image_path, "rb") as image_file:
                    encoded_image = base64.b64encode(image_file.read()).decode('utf-8')
                
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {self.api_key}"
                }
                
                payload = {
                    "model": self.api_model,
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "请仔细分析这张图片中的所有文字内容..."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{encoded_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    "max_tokens": 2000
                }
                
                response = requests.post(
                    self.api_url,
                    headers=headers,
                    json=payload,
                    timeout=60
                )
                
                response_data = response.json()
                if 'error' in response_data:
                    raise Exception(response_data['error']['message'])
                
                return response_data['choices'][0]['message']['content']
                
            except Exception as e:
                if retry_count < max_retries:
                    retry_count += 1
                    time.sleep(retry_delay)
                    retry_delay *= 1.5
                    continue
                raise Exception(f"处理图片时出错: {str(e)} (已重试 {retry_count} 次)")

ocr_processor = OCRProcessor()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有选择文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if file:
        filename = secure_filename(file.filename)
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            result = ocr_processor.process_image(filepath)
            return jsonify({'result': result})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)

@app.route('/export/docx', methods=['POST'])
def export_docx():
    try:
        data = request.json
        if not data or 'content' not in data:
            return jsonify({'error': '没有提供内容'}), 400
        
        content = data['content']
        
        # 创建Word文档
        doc = Document()
        doc.add_heading('OCR识别结果', 0)
        
        # 添加内容
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.startswith('=== 图片:'):
                # 添加图片标题
                doc.add_heading(para.strip(), level=1)
            else:
                # 添加正文内容
                doc.add_paragraph(para)
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # 发送文件
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f"OCR结果_{time.strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/export/txt', methods=['POST'])
def export_txt():
    try:
        data = request.json
        if not data or 'content' not in data:
            return jsonify({'error': '没有提供内容'}), 400
        
        content = data['content']
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
        with open(temp_file.name, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # 发送文件
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f"OCR结果_{time.strftime('%Y%m%d_%H%M%S')}.txt",
            mimetype='text/plain'
        )
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

@app.route('/config', methods=['GET', 'POST'])
def config():
    if request.method == 'POST':
        data = request.json
        ocr_processor.api_key = data.get('api_key', ocr_processor.api_key)
        ocr_processor.api_url = data.get('api_url', ocr_processor.api_url)
        ocr_processor.api_model = data.get('api_model', ocr_processor.api_model)
        ocr_processor.max_retries = int(data.get('max_retries', ocr_processor.max_retries))
        ocr_processor.save_config()
        return jsonify({'status': '配置已更新'})
    else:
        # 返回当前配置（不包含API密钥）
        return jsonify({
            'api_url': ocr_processor.api_url,
            'api_model': ocr_processor.api_model,
            'max_retries': ocr_processor.max_retries
        })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
